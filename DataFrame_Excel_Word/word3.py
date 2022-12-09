from docx import Document
from docx.shared import Inches


# Exemplo: https://python-docx.readthedocs.io/en/latest/

# Criando um doc word simples:

documento = Document()  # Inicializando

titulo_word = str(input('Informe o título no arquivo: '))

# documento.add_heading('Font Style: Roboto', 2)  # Estilo da fonte

documento.add_heading(f'{titulo_word}', 0)  # Título

paragrafo = str(input('\nInforme um parágrafo: '))

prgf = documento.add_paragraph(f'{paragrafo} ')  # Parágrafo

palavra = str(input('\nInforme uma palavra para ficar em negrito: '))
palavra2 = str(input('Informe uma palavra para ficar em itálico: '))

prgf.add_run(f'{palavra}').bold = True  # Adicionando uma palavra no parágrafo e também colocando em negrito
prgf.add_run(f' e {palavra2}.').italic = True  # Adicionando outra palavra e colocando em itálico

cabecalho = str(input('\nInforme o cabeçalho: '))

documento.add_heading(f'{cabecalho}', level=1)  # Cabeçalho

citacao = str(input('\nInforme uma citação: '))

documento.add_paragraph(f'{citacao}', style='Intense Quote')  # Citação e estilo

# Adicionando imagem (ela deve estar na mesma pasta do arquivo .py):

documento.add_picture('Python.svg.png', width=Inches(1.25))

# Podemos adicionar itens em uma lista ordenada = List Number || List bullet para lista desordenadas
# Neste caso, não usaremos, vamos aidicionar itens em uma lista com ordem de comprar, ID e produto

produto = str(input('\nInforme o produto: '))
produto2 = str(input('Informe outro produto: '))
produto3 = str(input('Informe mais um produto: '))


itens = (
    (1, '1000', f'{produto}'),
    (2, '1001', f'{produto2}'),
    (3, '1002', f'{produto3}')
)  # Adicionando itens em uma tabela

tabela = documento.add_table(rows=1, cols=3)  # Criando a tabela

# Configurando as células
celulas = tabela.rows[0].cells
celulas[0].text = 'Ordem de compra'
celulas[1].text = 'ID'
celulas[2].text = 'Produto'

# Looping para adicionar as células na tabela:

for ord_compra, ID, prdt in itens:
    row_celulas = tabela.add_row().cells
    row_celulas[0].text = str(ord_compra)
    row_celulas[1].text = ID
    row_celulas[2].text = prdt

documento.add_page_break()  # Finalizando

nome_arquivo = str(input('\nInforme o nome do arquivo: '))

documento.save(f'{nome_arquivo}.docx')  # Salvando
